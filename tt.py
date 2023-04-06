import os
import openai
import hashlib
import ebooklib
import csv
import requests
import json
from io import BytesIO
from ebooklib import epub
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH



# Load your API key from an environment variable or secret management service
openai.api_key = "API KEY HERE"
api_key = os.environ.get("API KEY HERE")
url = "https://api.openai.com/v1/images/generations"
globaldesc=""
#------------------------------------------------------------------------------------------------------------------------------------------------

def PR(prompt, system_content, temperature=0.5, token=12, pp=0.5, fp=0.5):
    messages = [{'role': 'system', 'content': system_content}, {'role': 'user', 'content': prompt}]
    response = openai.ChatCompletion.create(model="gpt-3.5-turbo", messages=messages, temperature=temperature, max_tokens=token, presence_penalty=pp,frequency_penalty=fp)
    system_message = response.choices[0].get('message', {}).get('content', '').strip()
    prompt_message = prompt.strip()
    #print("System message:", system_content)
    #print("Prompt message:", prompt_message)
    #print("Generated response:", system_message)
    return system_message

def APD(text, file_name):
    with open(file_name, "a+") as f:
        f.write(text + "\n")

def Title(text, place):
    return text[6:]

def txt_to_docx(txt_file, docx_file,town):
    # Create a new Document
    document = Document()
    # Set the page size and margins
    section = document.sections[0]
    section.page_width = Inches(4)
    section.page_height = Inches(6)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

     # Add title page
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run("Exploring " + town)
    title_run.font.size = Pt(18)
    title_run.bold = True
    title_paragraph.add_run("\n\nYour Guide to History, Culture, and Fun")
    document.add_page_break()

    # Add copyright page
    copyright_paragraph = document.add_paragraph()
    copyright_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    copyright_paragraph.add_run("Copyright 2023").font.size = Pt(8)

    # Set the position of the copyright notice to the bottom of the copyright page
    copyright_paragraph_format = copyright_paragraph.paragraph_format
    copyright_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    copyright_paragraph_format.space_after = Pt(0)
    document.add_page_break()

    # Add dedication page
    dedication_paragraph = document.add_paragraph()
    dedication_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dedication_paragraph.add_run("Dedication\n\n")
    dedication_paragraph.add_run(PR("Please write me 1 paragraph for the dedication of a history, culture, reference, and travel guide for " + town + ", considering it was written by one person and should be a generic dedication to travelers and explorers.", "You are a professional writer for Lonely Planet writing the dedication page of a book about the history, culture, reference, and travel guide for small town of " + town + ".Make sure to vary your words and grammar so phrases do not repeat too much.", 0.5, 1024)).italic = True
    document.add_page_break()

    # Read the input text file
    with open(txt_file, 'r') as f:
        txt_lines = f.readlines()
        txt_lines = [line for line in txt_lines if line.strip()]

    # Process the text file lines
    for line in txt_lines:
        if line.startswith('H1'):
            document.add_page_break()
            title = line[2:].strip()
            heading = document.add_heading(level=1)
            heading_run = heading.add_run(title)
            heading_run.font.size = Pt(18)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer API KEY HERE"
            }
            imgprompt = PR("I want you to write a DALL-E image generation prompt to generate an image related to " + title + " in the town of " + town, "You write image generation prompts for DALL-E from the given input request. For example, if you were asked to write a prompt for an image about the geography of Damascus, Virginia, you may output: Mountainous terrain with dense forests and trees, peaceful and serene, in the vicinity of Damascus, VA, USA. Shot on a Canon EOS R6 with a Canon RF 24-105mm f/4L IS USM Lens, 4K film still, natural lighting, vibrant colors, crisp details, and soft shadows.", 0.5, 2048)
            print("Generating Image For: " + title + " : " + town)
            data = {
                "prompt": imgprompt,
                "n": 1,
                "size": "1024x1024"
            } 

            response = requests.post(url, headers=headers, data=json.dumps(data))

            if response.status_code == 200:
                result = response.json()
                # Download the image from a link
                image_url = result['data'][0]['url']
                response = requests.get(image_url)
                img = BytesIO(response.content)
                # Add a paragraph with an image
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # center the paragraph
                run = paragraph.add_run()
                run.add_picture(img, width=Inches(3))  # adjust width as necessary

        elif line.startswith('H2'):
            title = line[2:].strip()
            heading = document.add_heading(level=2)
            heading_run = heading.add_run(title)
            heading_run.font.size = Pt(16)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        else:
            content = line.strip()
            paragraph = document.add_paragraph(content)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Save the generated DOCX file
    document.save(docx_file)

def read_towns_csv(csv_file):
    towns = []

    with open(csv_file, newline='') as f:
        reader = csv.reader(f)
        for row in reader:
            town, state = row[0].split(', ')
            towns.append([town, state])

    return towns
#------------------------------------------------------------------------------------------------------------------------------------------------

towns = read_towns_csv('towns.csv')

nouse=input();

for town_name, state in towns:
    #print(f"{town}, {state}")

    #town_name = input("Enter the town's name: ")
    #state = input("Enter the state: ")
    town = town_name + ", " + state
    savefile = town_name + " " + state + ".txt"

    book = [
        # Chapter 1
        [
            "H11.0 Introduction",
            "H21.1 Purpose of the book",
        ],
        # Chapter 2
        [
            "H12.0 Town's History",
            "H22.1 Early Settlement",
            "H22.2 Development and Growth",
            "H22.3 Notable Events and People",
            "H22.4 Modern Era",
        ],
        # Chapter 3
        [
            "H13.0 Geography and Environment",
            "H23.1 Physical Location",
            "H23.2 Topography",
            "H23.3 Climate",
            "H23.4 Natural Resources",
        ],
        # Chapter 4
        [
            "H14.0 Demographics",
            "H24.1 Population",
            "H24.2 Ethnicity and Culture",
            "H24.3 Education",
            "H24.4 Income and Employment",
        ],
        # Chapter 5
        [
            "H15.0 Governance and Politics",
            "H25.1 Local Government Structure",
            "H25.2 Key Officials and Representatives",
            "H25.3 Political Affiliations",
            "H25.4 Major Political Issues",
        ],
        # Chapter 6
        [
            "H16.0 Economy",
            "H26.1 Major Industries and Employers",
            "H26.2 Small Business Scene",
            "H26.3 Agriculture",
            "H26.4 Tourism",
        ],
        # Chapter 7
        [
            "H17.0 Community and Culture",
            "H27.1 Community Events and Traditions",
            "H27.2 Local Landmarks and Attractions",
            "H27.3 Arts and Entertainment",
            "H27.4 Sports and Recreation",
        ],
        # Chapter 8
        [
            "H18.0 Infrastructure and Services",
            "H28.1 Transportation",
            "H28.2 Healthcare",
            "H28.3 Education and Libraries",
            "H28.4 Utilities and Public Services",
        ],
        # Chapter 9
        [
            "H19.0 Conclusion",
            "H29.1 Reflection on the Town's Uniqueness and Character",
            "H29.2 Final Thoughts and Recommendations.",
        ]
    ]
    ShortBook = []

    for chapter in book:
        new_chapter = []
        for section in chapter:
            new_section = section[6:]
            new_chapter.append(new_section)
        ShortBook.append(new_chapter)

    #print(ShortBook)
    #print("PROGRESS:",end="")
    total_chars = 0
    total_words = 0
    total_sections = sum([len(c) for c in book])
    completed_sections = 0

    print("Title: Exploring "+town+": Your Guide to History, Culture, and Fun")

    for i, chapter in enumerate(book):
        for j, section in enumerate(chapter):
            print("Writing About The " + ShortBook[i][j] + " of " + town)
            if(j==0):
                prompt = "Write me a 1 paragraph brief introduction for a chapter about the " + ShortBook[i][j] + " of " + town + "."
            else:
                prompt = "Write me 8-10 extensive detailed and informative paragraphs about the " + ShortBook[i][j] + " of " + town + ". Only write strictly about " + ShortBook[i][j] + " and do not progess into any other topics/eras/areas/sections that lie outside of the " + ShortBook[i][j] + "."
            response = "\n\n" + book[i][j] + "\n" + PR(prompt, "You are a professional travel guide writer writing a book about the small town of " + town + ". You are asked to write a section for a book, and take care to not write any more than appropriate for the chapters topic, and not write past the strict topic. Only write strictly about " + ShortBook[i][j] + " and do not progess into any other topics/eras/areas/sections that lie outside of the " + ShortBook[i][j] + ". Remember, the town has already been introduced in earlier sections, so you don't have to give the reader an intro to the towns location or basic data, unless the chapter calls for it. For reference about the other topics so you know what to and not to write about, here is a list of the chapters in order: Purpose of the book,Early Settlement,Development and Growth,Notable Events and People,Modern Era,Physical Location,Topography,Climate,Natural Resources,Population,Ethnicity and Culture,Education,Income and Employment,Local Government Structure,Key Officials and Representatives,Political Affiliations,Major Political Issues,Major Industries and Employers,Small Business Scene,Agriculture,Tourism,Community Events and Traditions,Local Landmarks and Attractions,Arts and Entertainment,Sports and Recreation,Transportation,Healthcare,Education and Libraries,Utilities and Public Services,Reflection on the Town's Uniqueness and Character,Final Thoughts and Recommendations.Make sure to vary your words and grammar so phrases do not repeat too much.Make sure to only include factual informaation.", 0.5, 2048)
            num_chars = len(response)
            num_words = len(response.split(' '))
            total_chars += num_chars
            total_words += num_words
            completed_sections += 1
            print(f" {completed_sections}/{total_sections} sections completed ({completed_sections/total_sections*100:.2f}%). {num_chars} CHAR. {num_words} WORDS. Total: {total_chars} CHAR. {total_words} WORDS. Saving...")
            APD(response, savefile)

    print(f"All sections completed. {total_chars} CHAR. {total_words} WORDS saved!") 

    APD("\n\nH1References & Guides\n", savefile)

    ref_sec = [      
        ["H2A. Dining Guide",         "Create a detailed list of 20 recommended restaurants and cafes in the area, in order from closest to farthest away, including information on their cuisine and price range. Give me exact distance from town, driving distances, turns, exits, & driving directions.I don't need real time traffic, GPS, road conditions, or time of day/year, just directions from the towns  main street/road.."],
        ["H2B. Outdoor Activities Guide",         "Create a detailed list of 20 local parks, hiking trails, and other outdoor recreation opportunities, in order from closest to farthest away, along with details on equipment rental, guided tours, and other related services. Give me exact distance from town, driving distances, turns, exits, & driving directions.I don't need real time traffic, GPS, road conditions, or time of day/year, just directions from the towns  main street/road.."],
        ["H2C. Historical Sites and Museums",         "Create a detailed list of 20 historical sites, museums, and landmarks, in order from closest to farthest away that visitors may be interested in exploring. Give me exact distance from town, driving distances, turns, exits, & driving directions.I don't need real time traffic, GPS, road conditions, or time of day/year, just directions from the towns  main street/road.."],
        ["H2D. Local Artisan and Crafts Guide",         "Create a detailed list of 20 local artisans and crafters, in order from closest to farthest away, by providing a directory of shops, studios, local boutiques, antique stores, and specialty shops that sell unique and handmade items. Give me exact distance from town, driving distances, turns, exits, & driving directions.I don't need real time traffic, GPS, road conditions, or time of day/year, just directions from the towns  main street/road.."],
        ["H2E. Local Transportation Options",         "Provide detailed information on local transportation options, such as buses, taxis, and car rental services, to help visitors get around town easily. Please also list the nearest airports, in order from closest to farthest away."],
        ["H2F. Shopping Guide",         "Create a detailed list of 20 recommended large shopping destinations in the area, in order from closest to farthest away, such as major outlet stores, malls, and shopping centers (but not studios, local boutiques, antique stores, or specialty shops as these have already been covered) . Give me exact distance from town, driving distances, turns, exits, & driving directions.I don't need real time traffic, GPS, road conditions, or time of day/year, just directions from the towns main street/road."],
        ["H2G. Local Services Directory",         "Create a detailed directory of local services, such as pharmacies, hospitals, banks, police stations, city hall, and post offices, that visitors may need during their stay."],
        ["H2H. Local Wineries and Breweries",         "Create a detailed list of 20 local wineries and breweries, in order from closest to farthest away, providing information on tours, tastings, and other related experiences. Give me exact distance from town, driving distances, turns, exits, & driving directions.I don't need real time traffic, GPS, road conditions, or time of day/year, just directions from the towns  main street/road.."],
        ["H2I. Local Sports Teams",         "If there are any local sports teams in the area, include detailed information on their schedules, ticket prices, and other related details."],
        ["H2J. Unique Local Animals, Insects, & Wildlife",         "Create a detailed list of 10 Local Unique Animals, Insects, & Wildlife in the area, include information them. Provide name, type, scientific name, habitat, attributes, and a brief summary as well as hazards associated."],
        ["H2K. Unique Local Plants & Trees",         "Create a detailed list of 10 Local Unique Plants & Trees in the area, include information them. Provide name, type, scientific name, habitat, attributes, and a brief summary as well as hazards associated."],
        ["H2M. Local Legends",         "If there are any legends in the area, please include detailed information on the location, background, accessability to the site, and related information."],
        ["H2N. Notable Customs & Laws",         "If there are any customs or laws that differ to a considerable degree in comparison to other areas, please make note of them. Some exmaples, which you may or may not include as well as add upon, is: native american reserveations, dry counties, legalized cannabis, public smoking bans, blue laws, sales tax, traffic laws, gambling laws, and other details a tourist may not immediately know."]
    ]

    for title, prompt in ref_sec:
        print("Writing " + title[5:] + " For " + town, end="")
        response = "\n\n" + title + "\n" + PR(prompt, "You are a professional travel guide writer writing the extensive detailed reference section and tourism guide portion of a book about the small town of " + town + ".Make sure to vary your words and grammar so phrases do not repeat too much.Make sure to only include factual information (although this also includes any local legends, myths, cryptoids, or other culturly relevant legends).", 0.45, 2048, 0.05, 0.05)
        char_count = len(response)
        word_count = len(response.split())
        APD(response, savefile)
        print("          Saved. " + str(char_count) + " CHAR, " + str(word_count) + " WORDS") 
    outputfile = savefile.replace('.txt', '.docx')
    txt_to_docx(savefile,outputfile,town)

